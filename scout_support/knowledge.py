"""
Knowledge Hub — Document processing, chunking, embedding, and retrieval for Scout Support.

Contains:
- KnowledgeService: Upload processing, text extraction, chunking, embedding generation
- Similarity search for relevant knowledge retrieval during AI intake
- Ticket-based learning: extract resolution patterns from completed tickets
"""

import hashlib
import json
import logging
import math
import os
from datetime import datetime
from typing import Dict, List, Optional, Tuple

from openai import OpenAI

logger = logging.getLogger(__name__)

EMBEDDING_MODEL = "text-embedding-3-large"
EMBEDDING_DIMENSIONS = 3072
CHUNK_SIZE = 1500
CHUNK_OVERLAP = 200
MAX_KNOWLEDGE_RESULTS = 5
SIMILARITY_THRESHOLD = 0.30

KNOWLEDGE_CATEGORIES = {
    'sop': 'Standard Operating Procedure',
    'process_guide': 'Process Guide',
    'bullhorn_guide': 'Bullhorn Guide',
    'troubleshooting': 'Troubleshooting Guide',
    'policy': 'Company Policy',
    'training': 'Training Material',
    'resolution_pattern': 'Resolution Pattern (Learned)',
    'other': 'Other',
}


class KnowledgeService:

    def __init__(self):
        api_key = os.environ.get('OPENAI_API_KEY')
        self.openai_client = OpenAI(api_key=api_key) if api_key else None

    def process_uploaded_document(self, title: str, file_storage, category: str = 'other',
                                  description: str = '', uploaded_by: str = '') -> Optional['KnowledgeDocument']:
        from extensions import db
        from models import KnowledgeDocument

        filename = file_storage.filename or 'unknown'
        ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''

        file_bytes = file_storage.read()
        if not file_bytes:
            logger.warning(f"Empty file uploaded: {filename}")
            return None

        raw_text = self._extract_text(file_bytes, ext, filename)
        if not raw_text or len(raw_text.strip()) < 20:
            logger.warning(f"Insufficient text extracted from {filename} ({len(raw_text or '')} chars)")
            return None

        doc = KnowledgeDocument(
            title=title,
            filename=filename,
            doc_type='uploaded',
            category=category,
            description=description,
            raw_text=raw_text,
            status='processing',
            uploaded_by=uploaded_by,
        )
        db.session.add(doc)
        db.session.flush()

        chunks = self._chunk_text(raw_text)
        self._create_entries_with_embeddings(doc, chunks)

        doc.status = 'active'
        db.session.commit()
        logger.info(f"Processed document '{title}' ({filename}): {len(chunks)} chunks created")
        return doc

    def learn_from_ticket(self, ticket_id: int) -> Optional['KnowledgeDocument']:
        from extensions import db
        from models import SupportTicket, KnowledgeDocument

        ticket = SupportTicket.query.get(ticket_id)
        if not ticket or ticket.status not in ('completed', 'closed'):
            return None

        existing = KnowledgeDocument.query.filter_by(
            source_ticket_id=ticket_id, doc_type='ticket_resolution'
        ).first()
        if existing:
            logger.info(f"Knowledge already extracted from ticket {ticket.ticket_number}")
            return existing

        resolution_text = self._build_resolution_text(ticket)
        if not resolution_text or len(resolution_text.strip()) < 50:
            return None

        doc = KnowledgeDocument(
            title=f"Resolution: {ticket.subject[:200]}",
            doc_type='ticket_resolution',
            category='resolution_pattern',
            description=f"Learned from ticket {ticket.ticket_number} ({ticket.category})",
            source_ticket_id=ticket_id,
            raw_text=resolution_text,
            status='processing',
            uploaded_by='system',
        )
        db.session.add(doc)
        db.session.flush()

        chunks = self._chunk_text(resolution_text)
        self._create_entries_with_embeddings(doc, chunks)

        doc.status = 'active'
        db.session.commit()
        logger.info(f"Learned resolution from ticket {ticket.ticket_number}: {len(chunks)} chunks")
        return doc

    def retrieve_relevant_knowledge(self, query_text: str, top_k: int = MAX_KNOWLEDGE_RESULTS,
                                     threshold: float = SIMILARITY_THRESHOLD) -> List[Dict]:
        from models import KnowledgeEntry, KnowledgeDocument

        if not self.openai_client:
            return []

        query_embedding = self._generate_embedding(query_text)
        if not query_embedding:
            return []

        entries = KnowledgeEntry.query.join(KnowledgeDocument).filter(
            KnowledgeDocument.status == 'active',
            KnowledgeEntry.embedding_vector.isnot(None),
        ).all()

        if not entries:
            return []

        scored = []
        for entry in entries:
            try:
                entry_vec = json.loads(entry.embedding_vector)
                similarity = self._cosine_similarity(query_embedding, entry_vec)
                if similarity >= threshold:
                    scored.append((entry, similarity))
            except (json.JSONDecodeError, TypeError):
                continue

        scored.sort(key=lambda x: x[1], reverse=True)
        top_results = scored[:top_k]

        results = []
        for entry, score in top_results:
            doc = entry.document
            results.append({
                'content': entry.content,
                'title': doc.title,
                'category': doc.category,
                'doc_type': doc.doc_type,
                'similarity': round(score, 4),
                'entry_id': entry.id,
                'document_id': doc.id,
            })

        return results

    def build_knowledge_context(self, ticket_subject: str, ticket_description: str,
                                 ticket_category: str = '') -> str:
        query = f"{ticket_subject}\n{ticket_description}"
        if ticket_category:
            from scout_support_service import CATEGORY_LABELS
            label = CATEGORY_LABELS.get(ticket_category, ticket_category)
            query = f"[{label}] {query}"

        results = self.retrieve_relevant_knowledge(query)
        if not results:
            return ''

        sections = []
        for i, r in enumerate(results, 1):
            cat_label = KNOWLEDGE_CATEGORIES.get(r['category'], r['category'])
            sections.append(
                f"--- Knowledge Reference {i} (relevance: {r['similarity']:.0%}) ---\n"
                f"Source: {r['title']} [{cat_label}]\n"
                f"{r['content']}"
            )

        return (
            "\n\nRelevant Knowledge Base References:\n"
            "The following information from our knowledge base may help you resolve this issue more accurately. "
            "Use these references as additional context — they contain SOPs, past resolution patterns, and guides.\n\n"
            + "\n\n".join(sections)
        )

    def delete_document(self, document_id: int) -> bool:
        from extensions import db
        from models import KnowledgeDocument

        doc = KnowledgeDocument.query.get(document_id)
        if not doc:
            return False

        db.session.delete(doc)
        db.session.commit()
        logger.info(f"Deleted knowledge document {document_id}: {doc.title}")
        return True

    def get_stats(self) -> Dict:
        from models import KnowledgeDocument, KnowledgeEntry

        total_docs = KnowledgeDocument.query.filter_by(status='active').count()
        uploaded = KnowledgeDocument.query.filter_by(status='active', doc_type='uploaded').count()
        learned = KnowledgeDocument.query.filter_by(status='active', doc_type='ticket_resolution').count()
        total_entries = KnowledgeEntry.query.join(KnowledgeDocument).filter(
            KnowledgeDocument.status == 'active'
        ).count()

        return {
            'total_documents': total_docs,
            'uploaded_documents': uploaded,
            'learned_resolutions': learned,
            'total_entries': total_entries,
        }

    def _extract_text(self, file_bytes: bytes, ext: str, filename: str) -> str:
        if ext == 'txt':
            return file_bytes.decode('utf-8', errors='replace')
        elif ext == 'pdf':
            return self._extract_pdf_text(file_bytes, filename)
        elif ext in ('docx', 'doc'):
            return self._extract_docx_text(file_bytes, ext, filename)
        else:
            return file_bytes.decode('utf-8', errors='replace')

    def _extract_pdf_text(self, file_bytes: bytes, filename: str) -> str:
        text = ''
        try:
            import fitz
            doc = fitz.open(stream=file_bytes, filetype='pdf')
            for page in doc:
                text += page.get_text() + '\n'
            doc.close()
        except Exception as e:
            logger.warning(f"PyMuPDF failed for {filename}: {e}")
            try:
                import io
                from PyPDF2 import PdfReader
                reader = PdfReader(io.BytesIO(file_bytes))
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + '\n'
            except Exception as e2:
                logger.error(f"All PDF extraction failed for {filename}: {e2}")
        return text.strip()

    def _extract_docx_text(self, file_bytes: bytes, ext: str, filename: str) -> str:
        text = ''
        if ext == 'docx':
            try:
                import io
                from docx import Document
                doc = Document(io.BytesIO(file_bytes))
                text = '\n'.join(p.text for p in doc.paragraphs if p.text.strip())
            except Exception as e:
                logger.error(f"DOCX extraction failed for {filename}: {e}")
        elif ext == 'doc':
            tmp_path = None
            try:
                import tempfile
                import subprocess
                with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as tmp:
                    tmp.write(file_bytes)
                    tmp_path = tmp.name
                result = subprocess.run(['antiword', tmp_path], capture_output=True, text=True, timeout=30)
                text = result.stdout
            except FileNotFoundError:
                logger.warning(f"antiword not installed — cannot extract .doc file {filename}")
            except Exception as e:
                logger.error(f"DOC extraction failed for {filename}: {e}")
            finally:
                if tmp_path and os.path.exists(tmp_path):
                    os.unlink(tmp_path)
        return text.strip()

    def _chunk_text(self, text: str) -> List[str]:
        if len(text) <= CHUNK_SIZE:
            return [text]

        chunks = []
        start = 0
        while start < len(text):
            end = start + CHUNK_SIZE

            if end < len(text):
                break_point = text.rfind('\n\n', start, end)
                if break_point == -1 or break_point <= start:
                    break_point = text.rfind('\n', start + CHUNK_SIZE // 2, end)
                if break_point == -1 or break_point <= start:
                    break_point = text.rfind(' ', start + CHUNK_SIZE // 2, end)
                if break_point > start:
                    end = break_point + 1

            chunks.append(text[start:end].strip())
            start = max(start + 1, end - CHUNK_OVERLAP)

        chunks = [c for c in chunks if len(c) >= 20]
        return chunks

    def _create_entries_with_embeddings(self, doc, chunks: List[str]):
        from extensions import db
        from models import KnowledgeEntry

        for i, chunk in enumerate(chunks):
            content_hash = hashlib.sha256(chunk.encode()).hexdigest()
            embedding = self._generate_embedding(chunk)

            entry = KnowledgeEntry(
                document_id=doc.id,
                chunk_index=i,
                content=chunk,
                content_hash=content_hash,
                embedding_vector=json.dumps(embedding) if embedding else None,
                embedding_model=EMBEDDING_MODEL if embedding else None,
            )
            db.session.add(entry)

    def _generate_embedding(self, text: str) -> Optional[List[float]]:
        if not self.openai_client:
            return None
        try:
            text = text[:30000]
            response = self.openai_client.embeddings.create(
                model=EMBEDDING_MODEL,
                input=text,
            )
            return response.data[0].embedding
        except Exception as e:
            logger.error(f"Embedding generation failed: {e}")
            return None

    def _cosine_similarity(self, vec_a: List[float], vec_b: List[float]) -> float:
        if len(vec_a) != len(vec_b):
            return 0.0
        dot = sum(a * b for a, b in zip(vec_a, vec_b))
        norm_a = math.sqrt(sum(a * a for a in vec_a))
        norm_b = math.sqrt(sum(b * b for b in vec_b))
        if norm_a == 0 or norm_b == 0:
            return 0.0
        return dot / (norm_a * norm_b)

    def _build_resolution_text(self, ticket) -> str:
        from extensions import db as _db

        parts = [
            f"Issue Category: {ticket.category}",
            f"Subject: {ticket.subject}",
            f"Description: {ticket.description}",
        ]

        if ticket.ai_understanding:
            try:
                understanding = json.loads(ticket.ai_understanding)
                if understanding.get('understanding'):
                    parts.append(f"AI Analysis: {understanding['understanding']}")
                if understanding.get('resolution_approach'):
                    parts.append(f"Resolution Approach: {understanding['resolution_approach']}")
            except (json.JSONDecodeError, TypeError):
                pass

        if ticket.proposed_solution:
            try:
                solution = json.loads(ticket.proposed_solution)
                if solution.get('description_admin'):
                    parts.append(f"Technical Solution: {solution['description_admin']}")
                if solution.get('execution_steps'):
                    steps_str = json.dumps(solution['execution_steps'], indent=2)
                    parts.append(f"Execution Steps: {steps_str}")
            except (json.JSONDecodeError, TypeError):
                parts.append(f"Solution: {ticket.proposed_solution}")

        if ticket.resolution_note:
            parts.append(f"Resolution Note: {ticket.resolution_note}")

        conversations = ticket.conversations.order_by(_db.text('created_at ASC')).all()
        clarifications = [c for c in conversations if c.email_type in ('clarification', 'user_reply')]
        if clarifications:
            convo_parts = []
            for c in clarifications[:5]:
                convo_parts.append(f"[{c.direction}] {c.body[:500]}")
            parts.append(f"Key Clarifications:\n" + "\n".join(convo_parts))

        actions = ticket.actions.filter_by(success=True).all()
        if actions:
            action_summaries = [a.summary or f"{a.action_type} on {a.entity_type} {a.entity_id}" for a in actions]
            parts.append(f"Executed Actions: " + "; ".join(action_summaries))

        return "\n\n".join(parts)
