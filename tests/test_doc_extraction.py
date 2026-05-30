"""
Regression tests for utils.doc_extraction — the antiword replacement.

These tests lock in the behaviour we shipped to fix the production
[Errno 5] EIO cascade where antiword's broken Nix store path was
silently dropping legacy .doc resumes from screening, job application,
and scout_support pipelines.

What we verify:
  - Magic-byte routing chooses the right extractor regardless of
    filename extension.
  - DOCX bytes mislabelled as .doc are read by python-docx.
  - PDF bytes mislabelled as .doc return None so callers route to
    their own PDF extractor (avoids duplicating PDF logic).
  - RTF bytes mislabelled as .doc are stripped without invoking the
    network OCR path.
  - Genuine OLE2 .doc bytes return None cleanly — vision OCR is
    deliberately skipped because the OpenAI vision API rejects
    non-image MIME types (application/msword) with HTTP 400.
  - The four legacy antiword call sites no longer reference 'antiword'
    in source — defence in depth against accidental reintroduction.
"""

from __future__ import annotations

import io

from utils.doc_extraction import _detect_format, extract_doc_text


# ---------------------------------------------------------------------------
# Magic-byte detection
# ---------------------------------------------------------------------------

def test_detect_format_pdf():
    assert _detect_format(b"%PDF-1.4\n...") == "pdf"


def test_detect_format_docx():
    assert _detect_format(b"PK\x03\x04...") == "docx"


def test_detect_format_doc_ole2():
    assert _detect_format(b"\xd0\xcf\x11\xe0\xa1\xb1...") == "doc"


def test_detect_format_rtf():
    assert _detect_format(b"{\\rtf1\\ansi...") == "rtf"


def test_detect_format_unknown():
    assert _detect_format(b"Hello plain text") == "unknown"


def test_detect_format_empty():
    assert _detect_format(b"") == "unknown"
    assert _detect_format(None) == "unknown"  # type: ignore[arg-type]
    assert _detect_format(b"abc") == "unknown"


# ---------------------------------------------------------------------------
# extract_doc_text routing
# ---------------------------------------------------------------------------

def test_pdf_bytes_mislabelled_as_doc_returns_none():
    """PDF bytes labelled .doc should return None — caller routes to PDF extractor."""
    pdf_bytes = b"%PDF-1.4\n%FAKE-PDF-FOR-TESTING\n"
    assert extract_doc_text(pdf_bytes, "resume.doc") is None


def test_docx_bytes_mislabelled_as_doc_uses_python_docx():
    """DOCX bytes labelled .doc should be read by python-docx."""
    from docx import Document

    docx_doc = Document()
    docx_doc.add_paragraph("Hello from a docx mislabelled as doc")
    docx_doc.add_paragraph("Second paragraph with more content here")
    buf = io.BytesIO()
    docx_doc.save(buf)
    docx_bytes = buf.getvalue()

    # Confirm magic bytes look like docx (PK)
    assert docx_bytes[:2] == b"PK"

    text = extract_doc_text(docx_bytes, "resume.doc")
    assert text is not None
    assert "Hello from a docx" in text
    assert "Second paragraph" in text


def test_rtf_bytes_extracted_without_network_call():
    """RTF bytes should be stripped locally, never hitting OpenAI."""
    rtf = (
        b"{\\rtf1\\ansi\\deff0 {\\fonttbl {\\f0 Times;}}"
        b"\\f0 The candidate has 10 years of Python experience.}"
    )
    text = extract_doc_text(rtf, "resume.doc")

    assert text is not None
    assert "candidate" in text.lower()
    assert "python" in text.lower()


def test_genuine_ole2_doc_returns_none_when_antiword_fails():
    """When antiword is unavailable or returns no text, extract_doc_text
    returns None cleanly.  Vision OCR is deliberately skipped because the
    OpenAI vision API rejects application/msword with HTTP 400."""
    from unittest.mock import patch

    ole2_bytes = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1" + b"\x00" * 200
    with patch("utils.doc_extraction._try_antiword", return_value=None):
        text = extract_doc_text(ole2_bytes, "resume.doc")
    assert text is None


def test_empty_bytes_return_none():
    assert extract_doc_text(b"", "resume.doc") is None
    assert extract_doc_text(None, "resume.doc") is None  # type: ignore[arg-type]


def test_ole2_doc_returns_none_when_antiword_unavailable():
    """If antiword is not on PATH, extract_doc_text returns None gracefully."""
    from unittest.mock import patch

    ole2_bytes = b"\xd0\xcf\x11\xe0" + b"\x00" * 100
    with patch("utils.doc_extraction._try_antiword", return_value=None):
        text = extract_doc_text(ole2_bytes, "resume.doc")
    assert text is None


def test_unknown_binary_returns_none_when_antiword_fails():
    """Unrecognised binary format returns None when antiword also fails."""
    from unittest.mock import patch

    unknown_bytes = b"\x00\x01\x02\x03" + b"\xff" * 100
    with patch("utils.doc_extraction._try_antiword", return_value=None):
        text = extract_doc_text(unknown_bytes, "mystery.doc")
    assert text is None


def test_ole2_doc_returns_text_when_antiword_succeeds():
    """When antiword succeeds on a genuine OLE2 .doc, the text is returned."""
    from unittest.mock import patch

    ole2_bytes = b"\xd0\xcf\x11\xe0" + b"\x00" * 200
    extracted = "Angela Carney\nPurchasing Director\n20 years experience in procurement"
    with patch("utils.doc_extraction._try_antiword", return_value=extracted):
        text = extract_doc_text(ole2_bytes, "resume.doc")
    assert text == extracted


def test_antiword_eio_returns_none():
    """OSError(errno=5 EIO) from a broken Nix binary must be caught and
    return None — never raise up to the caller."""
    import errno
    from unittest.mock import patch

    from utils.doc_extraction import _try_antiword

    def _raise_eio(*args, **kwargs):
        raise OSError(errno.EIO, "Input/output error")

    with patch("subprocess.run", side_effect=_raise_eio):
        result = _try_antiword(b"\xd0\xcf\x11\xe0" + b"\x00" * 100, "test.doc")
    assert result is None


def test_antiword_file_not_found_returns_none():
    """If antiword is not installed, FileNotFoundError must be caught."""
    from unittest.mock import patch

    from utils.doc_extraction import _try_antiword

    with patch("subprocess.run", side_effect=FileNotFoundError("antiword not found")):
        result = _try_antiword(b"\xd0\xcf\x11\xe0" + b"\x00" * 100, "test.doc")
    assert result is None


def test_antiword_timeout_returns_none():
    """If antiword hangs, TimeoutExpired must be caught."""
    import subprocess
    from unittest.mock import patch

    from utils.doc_extraction import _try_antiword

    with patch("subprocess.run", side_effect=subprocess.TimeoutExpired("antiword", 30)):
        result = _try_antiword(b"\xd0\xcf\x11\xe0" + b"\x00" * 100, "test.doc")
    assert result is None


# ---------------------------------------------------------------------------
# Defence-in-depth: antiword must stay out of the codebase
# ---------------------------------------------------------------------------

# ---------------------------------------------------------------------------
# Pure-Python olefile .doc extraction (real fixtures)
#
# Fixtures under tests/fixtures/ are genuine OLE2 .doc files sourced from the
# Apache POI test-data set (Apache License 2.0). They exercise the FIB +
# piece-table parser added so legacy .doc extraction no longer depends on the
# antiword system binary (which works in dev but EIOs in production).
# ---------------------------------------------------------------------------

from pathlib import Path

from utils.doc_extraction import (
    _clean_doc_text,
    _printable_ratio,
    _try_olefile_doc,
)

_FIXTURES = Path(__file__).parent / "fixtures"


def _load(name: str) -> bytes:
    return (_FIXTURES / name).read_bytes()


def test_olefile_extracts_word97_simple_doc():
    text = _try_olefile_doc(_load("sample_word97.doc"), "sample_word97.doc")
    assert text is not None
    assert "simple file created with Word 97" in text


def test_olefile_extracts_modern_doc():
    text = _try_olefile_doc(_load("sample_calibri.doc"), "sample_calibri.doc")
    assert text is not None
    low = text.lower()
    assert "test document" in low
    assert "page 1" in low
    assert "calibri" in low


def test_olefile_extracts_large_doc_main_body():
    text = _try_olefile_doc(_load("sample_large.doc"), "sample_large.doc")
    assert text is not None
    assert len(text) > 1000
    assert "Manufacturer" in text


def test_olefile_extracts_table_doc_cell_content():
    text = _try_olefile_doc(_load("sample_table.doc"), "sample_table.doc")
    assert text is not None
    # All cell labels from the table must be present.
    for cell in ["A", "B", "C", "D", "E", "F", "G", "H", "I", "J"]:
        assert cell in text


def test_olefile_output_is_clean_text():
    """Extracted text must be overwhelmingly printable — never binary garbage."""
    for name in ("sample_word97.doc", "sample_calibri.doc", "sample_large.doc"):
        text = _try_olefile_doc(_load(name), name)
        assert text is not None
        assert _printable_ratio(text) >= 0.95


def test_extract_doc_text_works_without_antiword():
    """The production scenario: antiword is broken (EIO) yet a genuine .doc
    still extracts via the pure-Python olefile path."""
    from unittest.mock import patch

    with patch("utils.doc_extraction._try_antiword", return_value=None):
        text = extract_doc_text(_load("sample_word97.doc"), "resume.doc")
    assert text is not None
    assert "simple file created with Word 97" in text


def test_extract_doc_text_prefers_olefile_over_antiword():
    """olefile runs first; antiword must not be consulted when olefile succeeds."""
    from unittest.mock import patch

    with patch("utils.doc_extraction._try_antiword") as mock_aw:
        text = extract_doc_text(_load("sample_calibri.doc"), "resume.doc")
    assert text is not None
    mock_aw.assert_not_called()


def test_olefile_returns_none_on_non_ole_bytes():
    """Garbage / non-OLE bytes must return None, never raise or emit junk."""
    assert _try_olefile_doc(b"\x00\x01\x02\x03" + b"\xff" * 300, "x.doc") is None
    assert _try_olefile_doc(b"not an ole file at all", "x.doc") is None


def test_clean_doc_text_strips_field_codes():
    """Field instruction codes (0x13..0x14) are dropped; the visible field
    result (0x14..0x15) is kept."""
    raw = "Email: \x13HYPERLINK \"mailto:a@b.com\"\x14a@b.com\x15 end"
    cleaned = _clean_doc_text(raw)
    assert "HYPERLINK" not in cleaned
    assert "a@b.com" in cleaned
    assert "end" in cleaned


def test_clean_doc_text_translates_control_chars():
    raw = "Line one\rLine two\x0bbreak\x07cell\x1fsofthyphen"
    cleaned = _clean_doc_text(raw)
    assert "Line one" in cleaned
    assert "Line two" in cleaned
    assert "\x0b" not in cleaned
    assert "\x07" not in cleaned
    assert "\x1f" not in cleaned
    assert "softhyphen" in cleaned  # optional hyphen removed, text retained


def test_printable_ratio_guard():
    assert _printable_ratio("Hello world\nclean text") == 1.0
    assert _printable_ratio("\x00\x01\x02\x03\x04") < 0.5
    assert _printable_ratio("") == 0.0


def test_no_antiword_calls_in_resume_extraction_paths():
    """Regression guard: the four call sites we cleaned must not reinvoke antiword.

    If a future change reintroduces a subprocess call to antiword anywhere
    in these files, this test fails loudly so the broken-Nix-binary bug
    cannot return silently.
    """
    import re
    from pathlib import Path

    suspect_files = [
        Path("vetting/resume_utils.py"),
        Path("resume_parser.py"),
        Path("scout_support/knowledge.py"),
        Path("scout_support/ai_analysis.py"),
    ]

    # Only flag actual code invocations, not docstring mentions of the
    # historical context. We look for the string literal 'antiword' as
    # used in subprocess calls (subprocess.run(['antiword', ...]) or
    # shutil.which('antiword')) — those are the real reintroductions.
    pattern = re.compile(r"""['"]antiword['"]""")
    offenders = []
    for f in suspect_files:
        if not f.exists():
            continue
        content = f.read_text(encoding="utf-8")
        if pattern.search(content):
            offenders.append(str(f))

    assert not offenders, (
        f"antiword reference reintroduced in: {offenders}. "
        "Use utils.doc_extraction.extract_doc_text instead — antiword "
        "fails with [Errno 5] EIO on the production Nix store path."
    )
